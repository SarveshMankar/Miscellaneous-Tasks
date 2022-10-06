print("Processing!")

from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
import os
import send2trash

# Function to Resize the image
def cal(image1, pw, ph):
    img = Image.open(image1)
    width, height = img.size
    iw = 0.0104166667 * width
    ih = 0.0104166667 * height

    pr = ph / pw
    ir = ih / iw
    dw = iw - pw
    dh = ir * dw
    mw = pw
    mh = ih - dh

    if mw > pw or mh > ph:
        dh = ih - ph
        dw = dh / ir
        mh = ph
        mw = iw - dw

    return mw, mh


# Function for aligning to Center
def align():
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


# Making a Object of Document()
document = Document()

# Page Width and Height
mpw = 5.90551
mph = 8.66142

# Adding the Heading
head = document.add_paragraph().add_run("Your Images!")
head.font.size = Pt(28)
head.bold = True
head.underline = True
head.font.name = "Algerian"
align()

# Extensions
extensions = ["png", "jpg", "jpeg", "gif", "raw", "psd", "tiff", "ai", "jfif"]

# Getting All images which are in the Directory
aa = os.listdir()
a = []
for i in aa:
    if i.split(".")[-1].lower() in extensions:
        a.append(i)

# List for Files which were added
lst = []

# List for Files which were not added
notlst = []

# For the First Image, as the Heading has been added we need to reduce the Page Height
if a != []:
    try:
        name = document.add_paragraph().add_run(a[0])
        name.font.size = Pt(18)
        name.bold = True
        name.underline = True
        name.font.name = "Bahnschrift SemiBold"
        align()

        mw, mh = cal(a[0], mpw, mph - 1.7)
        document.add_picture(a[0], width=Inches(mw), height=Inches(mh))
        align()
        document.add_page_break()

        lst.append(a[0])
    except:
        # Adding the File name which could not be added
        notlst.append(a[0])

# Adding the Photos to the Word File
for i in range(1, len(a)):
    try:
        name = document.add_paragraph().add_run(a[i])
        name.font.size = Pt(18)
        name.bold = True
        name.underline = True
        name.font.name = "Bahnschrift SemiBold"
        align()

        mw, mh = cal(a[i], mpw, mph - 0.5)
        document.add_picture(a[i], width=Inches(mw), height=Inches(mh))
        align()
        document.add_page_break()

        lst.append(a[i])
    except:
        # Adding the File name which could not be added
        notlst.append(a[i])

if notlst == []:
    print("Your Work is Done! :-)")
else:
    print(
        "Your Work is Done! But Due to Some reasons couldn't add the Files mentioned below! "
    )
    print(notlst)

# For Deleting the Files
ans = input(
    'Do you want to Delete the files, which were added in Word file?\nIf Yes press "Y", if No press "N"\n'
)
if ans.lower() == "y":
    for i in lst:
        send2trash.send2trash(i)

# Saving the Document
document.save("images.docx")

print("Thank You! :-)")
